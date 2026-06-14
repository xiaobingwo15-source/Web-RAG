import csv
import io
import logging
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# MIME type constants for readability
_MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_MIME_PNG = "image/png"
_MIME_JPEG = "image/jpeg"
_MIME_TIFF = "image/tiff"
_MIME_WEBP = "image/webp"
_IMAGE_MIMES = {_MIME_PNG, _MIME_JPEG, _MIME_TIFF, _MIME_WEBP}

HEADER_REPEAT_ROWS = 20


@dataclass
class TextExtractionResult:
    text: str
    metadata: dict = field(default_factory=dict)


def extract_text(file_bytes: bytes, mime_type: str) -> str:
    if mime_type == "application/pdf":
        return _extract_pdf(file_bytes)
    elif mime_type in ("text/plain", "text/markdown"):
        return _extract_text(file_bytes)
    elif mime_type == "text/csv":
        return _extract_csv(file_bytes)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return _extract_excel(file_bytes)
    elif mime_type == _MIME_DOCX:
        return _extract_docx(file_bytes)
    elif mime_type in _IMAGE_MIMES:
        raise ValueError(
            "Image files require async OCR processing. Use extract_text_with_metadata() with use_ocr=True."
        )
    else:
        logger.warning(f"Unsupported mime type {mime_type}, treating as text")
        return _extract_text(file_bytes)


async def extract_text_with_metadata(
    file_bytes: bytes,
    mime_type: str,
    use_ocr: bool = False,
    pdf_parser_mode: str = "auto",
    filename: str | None = None,
) -> TextExtractionResult:
    if mime_type == "application/pdf":
        from app.services.pdf_parser import extract_pdf

        result = await extract_pdf(
            file_bytes,
            parser_mode=pdf_parser_mode,
            use_ocr=use_ocr,
            filename=filename,
        )
        return TextExtractionResult(text=result.text, metadata=result.metadata)
    elif mime_type == _MIME_DOCX:
        text, meta = _extract_docx_with_metadata(file_bytes)
        return TextExtractionResult(text=text, metadata=meta)
    elif mime_type in _IMAGE_MIMES:
        text, meta = await _extract_image_with_ocr(file_bytes, mime_type, filename)
        return TextExtractionResult(text=text, metadata=meta)
    return TextExtractionResult(text=extract_text(file_bytes, mime_type), metadata={})


async def extract_text_with_ocr(
    file_bytes: bytes,
    mime_type: str,
    use_ocr: bool = False,
    pdf_parser_mode: str = "auto",
) -> str:
    result = await extract_text_with_metadata(
        file_bytes,
        mime_type,
        use_ocr=use_ocr,
        pdf_parser_mode=pdf_parser_mode,
    )
    return result.text


def _extract_text(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def _extract_csv(file_bytes: bytes) -> str:
    text = _extract_text(file_bytes)
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ""

    header = rows[0]
    header_line = "\t".join(header)
    lines = [header_line]

    for i, row in enumerate(rows[1:], 1):
        lines.append("\t".join(row))
        if i % HEADER_REPEAT_ROWS == 0:
            lines.append("")
            lines.append(header_line)

    return "\n".join(lines)


def _extract_excel(file_bytes: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    all_text = []

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        all_text.append(f"=== Sheet: {sheet} ===")

        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(cell) if cell is not None else "" for cell in row])

        if not rows:
            continue

        header = rows[0]
        header_line = "\t".join(header)
        all_text.append(header_line)

        for i, row in enumerate(rows[1:], 1):
            all_text.append("\t".join(row))
            if i % HEADER_REPEAT_ROWS == 0:
                all_text.append("")
                all_text.append(header_line)

    wb.close()
    return "\n".join(all_text)


def _extract_pdf(file_bytes: bytes) -> str:
    from app.services.pdf_parser import extract_pdf

    import asyncio

    try:
        running_loop = asyncio.get_running_loop()
    except RuntimeError:
        running_loop = None

    if running_loop and running_loop.is_running():
        raise RuntimeError("_extract_pdf cannot be used inside a running event loop; use extract_text_with_metadata")

    return asyncio.run(extract_pdf(file_bytes, parser_mode="pypdfium")).text


def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file, preserving structure as markdown."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        # Map Word heading styles to markdown headings
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
                level = min(max(level, 1), 6)
            except ValueError:
                level = 1
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    # Extract tables as markdown
    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")

    return "\n".join(lines)


def _extract_docx_with_metadata(file_bytes: bytes) -> tuple[str, dict]:
    """Extract text and metadata from a DOCX file."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))

    # Count structural elements
    paragraph_count = len(doc.paragraphs)
    heading_count = sum(
        1 for p in doc.paragraphs
        if (p.style.name or "").lower().startswith("heading")
    )
    table_count = len(doc.tables)

    # Extract core properties
    meta: dict = {
        "docx_paragraph_count": paragraph_count,
        "docx_heading_count": heading_count,
        "docx_table_count": table_count,
    }

    try:
        props = doc.core_properties
        if props.title:
            meta["title"] = props.title
        if props.author:
            meta["author"] = props.author
        if props.created:
            meta["created"] = str(props.created)
        if props.modified:
            meta["modified"] = str(props.modified)
    except Exception:
        pass  # metadata is optional

    text = _extract_docx(file_bytes)
    return text, meta


async def _extract_image_with_ocr(
    file_bytes: bytes,
    mime_type: str,
    filename: str | None = None,
) -> tuple[str, dict]:
    """Extract text from an image using Gemini vision OCR."""
    import base64
    from app.services.ocr_service import ocr_image_with_gemini

    logger.info("Running OCR on image file (%s, %d bytes)", mime_type, len(file_bytes))

    try:
        text = await ocr_image_with_gemini(file_bytes, mime_type)
    except Exception as e:
        logger.error("Image OCR failed: %s", e)
        raise ValueError(f"Failed to extract text from image: {e}") from e

    if not text or not text.strip():
        raise ValueError("No text could be extracted from the image")

    meta = {
        "ocr_used": True,
        "source_mime_type": mime_type,
        "extracted_char_count": len(text),
    }
    if filename:
        meta["source_filename"] = filename

    return text, meta
